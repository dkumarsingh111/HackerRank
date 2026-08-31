"""Capstone_Project-HPPCS01.

Fine-tuned version:
- Gemma 3 is used for candidate-fact extraction and CV generation.
- Llama 3.2 is used as an independent CV reviewer.
- Candidate-specific CV content is never hard-coded.
- Source facts are preserved and obvious contact fields are verified deterministically.
- LLM output is constrained with JSON/plain-text instructions and bounded generation.
- Empty ACHIEVEMENTS/PROJECTS sections are omitted when the source has no content.
- Input resumes may be TXT, PDF, DOCX, DOC, RTF or MD.
- One controlled regeneration is performed only when the reviewer identifies a problem.
"""

import argparse
import difflib
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List

try:
    import requests
except ImportError:
    requests = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


"""
# Input document ingestion
"""

# SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".rtf", ".md"}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def extract_text_from_file(file_path: Path) -> str:
    """Read a supported resume format and normalize it to plain text."""
    ext = file_path.suffix.lower()

    if ext in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="replace")

    if ext == ".pdf":
        if PdfReader is None:
            raise RuntimeError("PDF support requires pypdf. Run: pip install pypdf")
        reader = PdfReader(str(file_path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        if not text.strip():
            raise RuntimeError(
                f"No text could be extracted from {file_path.name}. "
                "The PDF may be scanned/image-only; OCR support is required for that file."
            )
        return text

    if ext == ".docx":
        if Document is None:
            raise RuntimeError("DOCX support requires python-docx. Run: pip install python-docx")
        doc = Document(str(file_path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if ext == ".rtf":
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
        raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
        raw = re.sub(r"[{}]", " ", raw)
        return re.sub(r"\s+", " ", raw).strip()

    if ext == ".doc":
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError(
                "DOC support requires LibreOffice (soffice) in PATH."
            )
        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, str(file_path)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice failed for {file_path.name}: {result.stderr.strip()}")
            candidates = list(Path(tmp).glob("*.txt"))
            if not candidates:
                raise RuntimeError(f"No text output produced for {file_path.name}")
            return candidates[0].read_text(encoding="utf-8", errors="replace")

    raise ValueError(f"Unsupported input format: {file_path.suffix}")


def discover_input_files(input_dir: Path) -> List[Path]:
    """Discover supported resume documents, including PDF, DOC, DOCX and TXT."""
    candidates = [
        p for p in input_dir.iterdir()
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_EXTENSIONS
        and p.stem.lower().startswith("")
    ]
    # If the same candidate exists in multiple formats (for example profile_01.txt
    # and profile_01.pdf), process only one. PDF/DOCX are preferred over TXT so a
    # real resume document is not silently ignored during a multi-format test.
    priority = {".pdf": 0, ".docx": 1, ".doc": 2, ".rtf": 3, ".txt": 4, ".md": 5}
    selected = {}
    for p in candidates:
        key = p.stem.lower()
        if key not in selected or priority[p.suffix.lower()] < priority[selected[key].suffix.lower()]:
            selected[key] = p
    return sorted(selected.values(), key=lambda p: p.name.lower())


"""
# Ollama
"""

def call_ollama(
    model: str,
    prompt: str,
    host: str,
    timeout: int = 360,
    json_mode: bool = False,
    num_predict: int = 2000,
) -> str:
    """Call Ollama with bounded generation so one profile cannot hang indefinitely."""
    if requests is None:
        raise RuntimeError("Install requests first: pip install requests")

    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "keep_alive": "5m",
        "options": {
            "temperature": 0.1,
            "top_p": 0.9,
            "num_predict": num_predict,
        },
    }

    # Ollama supports JSON-constrained output with format=json.
    if json_mode:
        payload["format"] = "json"

    response = requests.post(
        f"{host.rstrip('/')}/api/generate",
        json=payload,
        timeout=timeout,
    )
    response.raise_for_status()

    result = response.json()
    text = result.get("response", "")
    if not text.strip():
        raise ValueError(f"Ollama returned an empty response for model {model}.")
    return text.strip()


def extract_json(text: str) -> Dict[str, Any]:
    """Parse JSON even when a small model accidentally adds a markdown fence."""
    text = text.strip()

    try:
        value = json.loads(text)
        if isinstance(value, dict):
            return value
    except json.JSONDecodeError:
        pass

    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text).strip()

    try:
        value = json.loads(text)
        if isinstance(value, dict):
            return value
    except json.JSONDecodeError:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        value = json.loads(text[start:end + 1])
        if isinstance(value, dict):
            return value

    raise ValueError("Model did not return a valid JSON object.")


"""
# Source-grounding helpers
"""

def first_email(raw: str) -> str:
    match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", raw)
    return match.group(0) if match else ""


def first_phone(raw: str) -> str:
    # Handles examples such as 555-1234 and +91-90000-10002.
    matches = re.findall(r"(?<!\w)(?:\+?\d[\d\s().-]{6,}\d)(?!\w)", raw)
    return matches[0].strip() if matches else ""


def extract_name_from_source(raw: str) -> str:
    # Labeled profiles: "Name: Priya Menon"
    match = re.search(r"(?im)^\s*name\s*:\s*([^\n,]+)", raw)
    if match:
        return match.group(1).strip()

    # Free-form profile: "John Doe, phone ..."
    match = re.match(
        r"\s*([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z.'-]+){1,3})\s*,",
        raw,
    )
    return match.group(1).strip() if match else ""


def extract_location_from_source(raw: str) -> str:
    # Labeled contact field.
    match = re.search(r"(?im)^\s*contact\s*:\s*.*,\s*([^,\n]+)\s*$", raw)
    if match:
        return match.group(1).strip()

    # Explicit location/city labels.
    match = re.search(r"(?im)^\s*(?:location|city)\s*:\s*([^\n]+)", raw)
    return match.group(1).strip() if match else ""


def clean_string(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def clean_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        value = [value]
    if not isinstance(value, list):
        return []
    result = []
    for item in value:
        text = clean_string(item)
        if text and text not in result:
            result.append(text)
    return result


def normalize_profile(profile: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize model JSON into the schema used by the rest of the pipeline."""
    contact = profile.get("contact")
    if not isinstance(contact, dict):
        contact = {}

    normalized = {
        "name": clean_string(profile.get("name")),
        "contact": {
            "email": clean_string(contact.get("email")),
            "phone": clean_string(contact.get("phone")),
            "location": clean_string(contact.get("location")),
        },
        "summary": clean_string(profile.get("summary")),
        "education": clean_list(profile.get("education")),
        "experience": clean_list(profile.get("experience")),
        "skills": clean_list(profile.get("skills")),
        "achievements": clean_list(profile.get("achievements")),
        "projects": clean_list(profile.get("projects")),
    }
    return normalized


def merge_obvious_source_facts(profile: Dict[str, Any], raw: str) -> Dict[str, Any]:
    """Protect obvious source facts from a weak extraction model."""
    profile = normalize_profile(profile)

    source_name = extract_name_from_source(raw)
    source_email = first_email(raw)
    source_phone = first_phone(raw)
    source_location = extract_location_from_source(raw)

    if source_name:
        profile["name"] = source_name
    if source_email:
        profile["contact"]["email"] = source_email
    if source_phone:
        profile["contact"]["phone"] = source_phone
    if source_location:
        profile["contact"]["location"] = source_location

    return profile


def fallback_extract(raw: str) -> Dict[str, Any]:
    """Safe fallback when Gemma is unavailable or times out."""
    lines = [x.strip() for x in raw.splitlines() if x.strip()]

    name = extract_name_from_source(raw) or "Candidate"
    email = first_email(raw)
    phone = first_phone(raw)
    location = extract_location_from_source(raw)

    summary = ""
    for line in lines:
        if re.match(r"(?i)^(profile|summary|professional background)\s*:", line):
            summary = line.split(":", 1)[1].strip()
            break

    skills = []
    for line in lines:
        if re.match(r"(?i)^skills\s*:", line):
            skills = [x.strip() for x in line.split(":", 1)[1].split(",")]
            break

    education = []
    achievements = []
    projects = []

    for line in lines:
        low = line.lower()
        if low.startswith("education:"):
            education.append(line.split(":", 1)[1].strip())
        elif low.startswith("achievements:"):
            achievements.append(line.split(":", 1)[1].strip())
        elif low.startswith("projects:"):
            projects.append(line.split(":", 1)[1].strip())

    # Keep complete experience sentences as source evidence in fallback mode.
    experience = []
    for line in lines:
        low = line.lower()
        if any(
            marker in low
            for marker in (
                "worked at ",
                "current employer:",
                "previous:",
                "experience:",
            )
        ):
            experience.append(line)

    return normalize_profile(
        {
            "name": name,
            "contact": {
                "email": email,
                "phone": phone,
                "location": location,
            },
            "summary": summary,
            "education": education,
            "experience": experience,
            "skills": skills,
            "achievements": achievements,
            "projects": projects,
        }
    )


"""
# Job Description parsing
"""

def normalize_job_description(value: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize LLM output to the required downstream JD JSON schema."""
    return {
        "job_title": clean_string(value.get("job_title")),
        "requirements": clean_list(value.get("requirements")),
        "responsibilities": clean_list(value.get("responsibilities")),
        "keywords": clean_list(value.get("keywords")),
    }


def fallback_job_description(text: str) -> Dict[str, Any]:
    """Conservative fallback if Gemma cannot parse the job description."""
    lines = [x.strip(" •*-\t") for x in text.splitlines() if x.strip()]
    title = ""
    requirements: List[str] = []
    responsibilities: List[str] = []
    keywords: List[str] = []

    for line in lines:
        low = line.lower()
        if not title and re.search(r"\b(job title|position|role)\s*:", low):
            title = line.split(":", 1)[1].strip()
        if re.match(r"(?i)^(requirements?|qualifications?)\s*:", line):
            requirements.append(line.split(":", 1)[1].strip())
        elif re.match(r"(?i)^(responsibilities|duties)\s*:", line):
            responsibilities.append(line.split(":", 1)[1].strip())

    # Do not invent keywords. Preserve explicit technical/professional terms only
    # when they occur verbatim in the JD text.
    known_terms = [
        "Python", "SQL", "Java", "C++", "C#", "JavaScript", "TypeScript",
        "PyTorch", "TensorFlow", "LangChain", "LlamaIndex", "Hugging Face",
        "RAG", "LLM", "Machine Learning", "Deep Learning", "NLP", "Power BI",
        "Excel", "Pandas", "AWS", "Azure", "GCP", "Docker", "Kubernetes",
    ]
    low_text = text.lower()
    for term in known_terms:
        if term.lower() in low_text:
            keywords.append(term)

    return normalize_job_description({
        "job_title": title,
        "requirements": requirements,
        "responsibilities": responsibilities,
        "keywords": keywords,
    })


def parse_job_description(
    jd_text: str,
    model: str,
    host: str,
    timeout: int,
) -> Dict[str, Any]:
    """Use Gemma to extract JD requirements, responsibilities and keywords."""
    prompt = f"""You are a strict job-description information extraction system.

Extract ONLY information explicitly present in JOB DESCRIPTION.
Do not infer, improve, invent, or add requirements that are not stated.

Return JSON only in exactly this structure:

{{
  "job_title": "",
  "requirements": [],
  "responsibilities": [],
  "keywords": []
}}

Rules:

- job_title: copy the stated job/position title.
- requirements: extract explicit qualifications, education, experience,
  skills, technologies, and capabilities required by the employer.
- responsibilities: extract explicit duties and responsibilities.
- keywords: extract important technical, professional, domain, tool,
  framework, database, programming language, and AI/ML terms explicitly
  present in the JD.
- Do not create candidate qualifications.
- Do not add keywords that are not present in the JD.
- Do not use information from the candidate resume.
- Return JSON only.
- Do not return Markdown or explanatory text.

JOB DESCRIPTION:
{jd_text}
"""


    try:
        result = call_ollama(
            model,
            prompt,
            host,
            timeout=timeout,
            json_mode=True,
            num_predict=1800,
        )
        return normalize_job_description(extract_json(result))
    except Exception as exc:
        print(f"[WARN] Job description extraction fallback: {exc}")
        return fallback_job_description(jd_text)


def load_job_description(input_dir: Path, args: argparse.Namespace) -> str:
    """Load JD from file/argument or accept copy-pasted text interactively."""
    jd_file = input_dir / "job_description.txt"
    if getattr(args, "jd_file", ""):
        jd_file = Path(args.jd_file)

    if jd_file.exists() and jd_file.is_file():
        text = jd_file.read_text(encoding="utf-8", errors="replace").strip()
        if text:
            print(f"[INPUT] Using job description: {jd_file.name}")
            return text

    if getattr(args, "jd", ""):
        return args.jd.strip()

    print("[INPUT] No job_description.txt found.")
    print("[INPUT] Paste the job description below.")
    print("[INPUT] Type END on a new line when finished.")
    lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip().upper() == "END":
            break
        lines.append(line)
    text = "\n".join(lines).strip()
    return text


def save_job_description(jd: Dict[str, Any], output_dir: Path) -> None:
    """Persist the structured JD required for downstream resume tailoring."""
    (output_dir / "job_description.json").write_text(
        json.dumps(jd, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


"""
# Gemma: extraction
"""

def extract_profile(
    raw: str,
    model: str,
    host: str,
    timeout: int,
) -> Dict[str, Any]:
    """Use Gemma to structure facts without inventing or improving them."""
    prompt = f"""You are a strict information extraction system.

Extract ONLY information explicitly present in SOURCE PROFILE.
Do not infer, improve, paraphrase, or invent facts.

Return JSON only:
{{
  "name": "",
  "contact": {{"email": "", "phone": "", "location": ""}},
  "summary": "",
  "education": [],
  "experience": [],
  "skills": [],
  "achievements": [],
  "projects": []
}}

Rules:
- Copy names, employers, dates, skills, education, achievements and projects from the source.
- If a field is not present, use "" or [].
- Keep experience as source-supported statements.
- Do NOT create job titles that are not stated.
- Do NOT create metrics or accomplishments.
- Do NOT add technologies, degrees or responsibilities that are absent.
- Return JSON only. No markdown.

SOURCE PROFILE:
{raw}
"""

    try:
        result = call_ollama(
            model,
            prompt,
            host,
            timeout=timeout,
            json_mode=True,
            num_predict=2000,
        )
        profile = extract_json(result)
        return merge_obvious_source_facts(profile, raw)
    except Exception as exc:
        print(f"[WARN] Extraction fallback: {exc}")
        return fallback_extract(raw)


"""
# Gemma: CV generation
"""

def generation_prompt(
    profile: Dict[str, Any],
    raw: str,
    job_description: Dict[str, Any] | None = None,
    regeneration_reason: str = "",
) -> str:
    reason = ""
    if regeneration_reason:
        reason = f"""
A previous draft failed review for these reasons:
{regeneration_reason}

Correct those issues while keeping every statement source-grounded.
"""

    job_description = job_description or {}
    return f"""You are a professional resume writer.

Create an ATS-friendly resume using ONLY the candidate information provided in
the SOURCE PROFILE and STRUCTURED FACTS.

The TARGET JOB DESCRIPTION is used ONLY to identify relevant wording and
keywords. Never add a skill, technology, responsibility, achievement,
certification, education, employer, job title, date, metric, or experience
unless it is supported by the candidate source.

TARGET JOB DESCRIPTION:
{json.dumps(job_description, indent=2, ensure_ascii=False)}

SOURCE PROFILE:
{raw}

STRUCTURED FACTS:
{json.dumps(profile, indent=2, ensure_ascii=False)}

{reason}

IMPORTANT OUTPUT REQUIREMENTS:

1. Do NOT copy the SOURCE PROFILE as a paragraph.
2. Convert the source information into a structured professional resume.
3. Return ONLY the resume content.
4. Do NOT provide explanations, analysis, comments, notes, rationale,
   introductions, conclusions, or statements about the resume.
5. Do NOT start with phrases such as:
   "Okay, here's a revised resume"
   "Here is the resume"
   "I've created"
   "I've revised"
   "This resume"
6. Do NOT include "Source Profile", "Analysis", "Key Changes & Rationale",
   "Note", or similar explanatory sections.
7. Do NOT use Markdown formatting, code fences, JSON, or XML.
8. Use the following section headings when applicable:

PROFESSIONAL SUMMARY
SKILLS
PROFESSIONAL EXPERIENCE
EDUCATION
ACHIEVEMENTS
PROJECTS

9. Include a section ONLY when the candidate source contains information
   for that section.
10. Do NOT write "None", "N/A", "Not provided", or placeholders for empty
    sections.
11. Preserve all available candidate contact information.
12. Organize work experience by employer, job title, and dates when available.
13. Convert source responsibilities into concise professional resume bullets.
14. You may improve grammar and professional wording, but the meaning must
    remain supported by the source.
15. You may align wording with the target JD only when the candidate source
    supports that wording.
16. Target-job keywords that are not supported by the candidate source must
    NOT be inserted into the resume.
17. Do not create achievements from ordinary responsibilities.
18. Do not create projects when none are present.
19. Do not create certifications, technologies, degrees, or qualifications
    that are not present in the source.
20. Keep the resume concise and suitable for ATS processing.

GROUNDING EXAMPLES:

Allowed:
"talked to clients" -> "Client communication"

Allowed:
"cash register work" -> "Cash register operations"

Allowed:
"updated the old computer system" -> "Updated the computer system"

Not allowed:
"talked to clients" -> "Built long-term client relationships"

Not allowed:
"worked in sales" -> "Exceeded sales targets"

Not allowed:
"updated the computer system" -> "Implemented an enterprise technology
platform"

FINAL OUTPUT FORMAT EXAMPLE:

<CANDIDATE NAME>
<EMAIL> | <PHONE> | <LOCATION>

PROFESSIONAL SUMMARY
Professional summary based only on the source information.

SKILLS
• Skill from source
• Skill from source

PROFESSIONAL EXPERIENCE

<EMPLOYER> | <JOB TITLE> | <DATES>
• Responsibility supported by source

EDUCATION

<Degree or education item from source>

Return ONLY the final resume in this structure.
"""


def generate_cv(
    profile: Dict[str, Any],
    raw: str,
    model: str,
    host: str,
    timeout: int,
    regeneration_reason: str = "",
    job_description: Dict[str, Any] | None = None,
) -> str:
    """Generate a grounded CV with bounded output."""
    prompt = generation_prompt(profile, raw, job_description, regeneration_reason)

    try:
        cv = call_ollama(
            model,
            prompt,
            host,
            timeout=timeout,
            json_mode=False,
            num_predict=2000,
        )
        return clean_cv_text(cv)
    except Exception as exc:
        print(f"[WARN] Generation fallback: {exc}")
        return deterministic_cv(profile)


def clean_cv_text(text: str) -> str:
    """Remove model wrappers/placeholders without changing candidate facts."""
    text = text.strip()
    text = re.sub(r"^```(?:text|markdown)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text).strip()

    # Remove common model chatter.
    text = re.sub(
        r"^(here(?:'s| is) (?:a|the) .*?(?:resume|cv).*?:)\s*",
        "",
        text,
        flags=re.I | re.S,
    )

    # Remove Markdown emphasis around headings/bullets.
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"(?m)^\s*[-*]\s*", "• ", text)

    # Remove obvious placeholder lines.
    lines = []
    placeholder_patterns = (
        r"\[.*(?:specify|location|include|add|available).*?\]",
        r"^\s*n/?a\s*$",
        r"^\s*none\s*$",
    )
    for line in text.splitlines():
        if any(re.search(p, line, flags=re.I) for p in placeholder_patterns):
            continue
        lines.append(line.rstrip())

    return "\n".join(lines).strip()


"""
# Deterministic fallback
"""

def deterministic_cv(profile: Dict[str, Any]) -> str:
    """Fallback CV using only already-extracted candidate facts."""
    c = profile.get("contact", {})

    lines = [
        profile.get("name") or "Candidate",
        " | ".join(
            x
            for x in (
                c.get("email", ""),
                c.get("phone", ""),
                c.get("location", ""),
            )
            if x
        ),
    ]

    if profile.get("summary"):
        lines += ["", "PROFESSIONAL SUMMARY", profile["summary"]]

    if profile.get("skills"):
        lines += ["", "SKILLS"]
        lines += [f"• {x}" for x in profile["skills"]]

    if profile.get("experience"):
        lines += ["", "PROFESSIONAL EXPERIENCE"]
        lines += [f"• {x}" for x in profile["experience"]]

    if profile.get("education"):
        lines += ["", "EDUCATION"]
        lines += [f"• {x}" for x in profile["education"]]

    if profile.get("achievements"):
        lines += ["", "ACHIEVEMENTS"]
        lines += [f"• {x}" for x in profile["achievements"]]

    if profile.get("projects"):
        lines += ["", "PROJECTS"]
        lines += [f"• {x}" for x in profile["projects"]]

    return "\n".join(lines)


"""
# Independent Llama review
"""

def review_cv(
    cv: str,
    profile: Dict[str, Any],
    raw: str,
    model: str,
    host: str,
    timeout: int,
    job_description: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    """Use Llama for independent review."""
    job_description = job_description or {}
    prompt = f"""You are an independent resume quality reviewer.

Compare the generated CV against the SOURCE PROFILE and TARGET JOB DESCRIPTION.

Return JSON only:
{{
  "score": 0,
  "keyword_coverage": 0,
  "factual_consistency": 0,
  "format_score": 0,
  "missing_keywords": [],
  "unsupported_claims": [],
  "issues": [],
  "recommendation": "PASS or REGENERATE"
}}

Scoring:
- keyword_coverage: how well the CV covers relevant target-job keywords while preserving source facts.
- factual_consistency: whether every factual claim is supported by the source.
- format_score: readability and ATS structure.
- score: overall score.

Important:
- Missing phone, email, employer, date, education, skill, achievement or project
  that exists in SOURCE PROFILE is an issue.
- Unsupported achievements, metrics, responsibilities, job titles or skills are
  an issue.
- A target-job keyword is not a candidate skill unless SOURCE PROFILE supports it.
- Empty sections or placeholders are an issue.
- Do not give credit for information that is not in the source.
- List every material unsupported candidate claim in unsupported_claims.
- If factual consistency is below 90 OR important source information is missing
  OR unsupported claims exist, recommendation must be REGENERATE.
- A score of 0 cannot be PASS.
- Return JSON only.

SOURCE PROFILE:
{raw}

STRUCTURED FACTS:
{json.dumps(profile, indent=2)}

TARGET JOB DESCRIPTION:
{json.dumps(job_description, indent=2, ensure_ascii=False)}

GENERATED CV:
{cv}
"""

    try:
        result = extract_json(
            call_ollama(
                model,
                prompt,
                host,
                timeout=timeout,
                json_mode=True,
                num_predict=2000,
            )
        )
        return normalize_review(result)
    except Exception as exc:
        print(f"[WARN] Reviewer fallback: {exc}")
        return deterministic_review(cv, profile, raw, str(exc))


def normalize_review(review: Dict[str, Any]) -> Dict[str, Any]:
    """Make reviewer output safe and ensure PASS cannot have score 0."""
    def score(key: str, default: int = 0) -> int:
        try:
            value = int(float(review.get(key, default)))
        except (TypeError, ValueError):
            value = default
        return max(0, min(100, value))

    result = {
        "score": score("score"),
        "keyword_coverage": score("keyword_coverage"),
        "factual_consistency": score("factual_consistency"),
        "format_score": score("format_score"),
        "missing_keywords": clean_list(review.get("missing_keywords")),
        "unsupported_claims": clean_list(review.get("unsupported_claims")),
        "issues": clean_list(review.get("issues")),
        "recommendation": clean_string(review.get("recommendation")).upper(),
    }

    if result["recommendation"] not in {"PASS", "REGENERATE"}:
        result["recommendation"] = (
            "PASS"
            if result["score"] >= 90 and result["factual_consistency"] >= 90
            else "REGENERATE"
        )

    if result["score"] <= 0 or result["factual_consistency"] < 90:
        result["recommendation"] = "REGENERATE"

    return result


def deterministic_review(
    cv: str,
    profile: Dict[str, Any],
    raw: str,
    error_message: str,
) -> Dict[str, Any]:
    """Safety-net review when Llama cannot be reached."""
    low = cv.lower()
    missing = []

    important_values = [
        profile.get("name", ""),
        profile.get("contact", {}).get("email", ""),
        profile.get("contact", {}).get("phone", ""),
        profile.get("contact", {}).get("location", ""),
    ]

    for value in important_values:
        if value and value.lower() not in low:
            missing.append(value)

    for skill in profile.get("skills", []):
        if skill and skill.lower() not in low:
            missing.append(skill)

    coverage_items = [x for x in important_values + profile.get("skills", []) if x]
    coverage = round(
        100 * (len(coverage_items) - len(set(missing))) / max(1, len(coverage_items))
    )

    # Reviewer fallback must be conservative.
    score = min(coverage, 80)

    return {
        "score": score,
        "keyword_coverage": coverage,
        "factual_consistency": 0,
        "format_score": 80,
        "missing_keywords": sorted(set(missing)),
        "issues": [f"Independent reviewer unavailable: {error_message}"],
        "recommendation": "REGENERATE",
    }


"""
# Deterministic grounding, source-preservation and ATS guards
"""

def normalized_alnum(text: str) -> str:
    """Normalize text for punctuation-insensitive comparisons."""
    return re.sub(r"[^a-z0-9]+", " ", clean_string(text).lower()).strip()


def phrase_present(cv: str, phrase: str) -> bool:
    """Check whether a source phrase survives harmless formatting changes."""
    a = normalized_alnum(cv)
    b = normalized_alnum(phrase)
    if not b:
        return True
    if b in a:
        return True
    tokens = [t for t in b.split() if len(t) >= 3]
    if not tokens:
        return True
    return sum(token in a for token in tokens) / len(tokens) >= 0.70


def source_supported_target_keywords(
    raw: str, job_description: Dict[str, Any], evidence_text: str = ""
) -> Dict[str, Any]:
    """Identify JD keywords supported by the original source or explicit user evidence."""
    evidence = "\n".join(x for x in (raw, evidence_text) if x)
    supported = []
    unsupported = []
    for keyword in clean_list(job_description.get("keywords")):
        if phrase_present(evidence, keyword):
            supported.append(keyword)
        else:
            unsupported.append(keyword)
    return {
        "source_supported_keywords": supported,
        "source_unsupported_keywords": unsupported,
    }


def deterministic_grounding_check(
    cv: str,
    raw: str,
    profile: Dict[str, Any],
    job_description: Dict[str, Any],
    candidate_evidence: str = "",
) -> Dict[str, Any]:
    """Conservative post-generation guard against omissions and hallucinations."""
    issues: List[str] = []
    missing_source: List[str] = []
    unsupported_keywords: List[str] = []

    contact = profile.get("contact", {})
    for label, value in (
        ("name", profile.get("name", "")),
        ("email", contact.get("email", "")),
        ("phone", contact.get("phone", "")),
        ("location", contact.get("location", "")),
    ):
        if value and not phrase_present(cv, value):
            missing_source.append(f"{label}: {value}")

    for skill in profile.get("skills", []):
        if skill and not phrase_present(cv, skill):
            missing_source.append(f"skill: {skill}")

    cv_normalized = normalized_alnum(cv)
    for field in ("education", "experience", "achievements", "projects"):
        for item in profile.get(field, []):
            item = clean_string(item)
            if not item:
                continue
            if len(item.split()) <= 8:
                present = phrase_present(cv, item)
            else:
                tokens = list(dict.fromkeys(
                    t for t in re.findall(r"[A-Za-z0-9]+", item.lower())
                    if len(t) >= 5
                ))[:8]
                present = (
                    not tokens
                    or sum(t in cv_normalized for t in tokens) / len(tokens) >= 0.50
                )
            if not present:
                missing_source.append(f"{field}: {item}")

    support = source_supported_target_keywords(raw, job_description, candidate_evidence)
    for keyword in support["source_unsupported_keywords"]:
        if normalized_alnum(keyword) and normalized_alnum(keyword) in cv_normalized:
            unsupported_keywords.append(keyword)

    if missing_source:
        issues.append(
            "Source information was omitted or materially changed: "
            + "; ".join(missing_source[:10])
        )
    if unsupported_keywords:
        issues.append(
            "Unsupported target-job keywords were added to the CV: "
            + ", ".join(unsupported_keywords[:10])
        )

    if re.search(r"(?im)^\s*(?:n/?a|tbd|to be added|insert here)\s*$", cv):
        issues.append("The generated CV contains a placeholder.")
    if re.search(r"\[[^\]]+\]", cv):
        issues.append("The generated CV contains bracketed placeholder content.")

    if not profile.get("achievements") and re.search(
        r"(?im)^\s*ACHIEVEMENTS\s*$", cv
    ):
        issues.append("ACHIEVEMENTS section was added although the source has none.")
    if not profile.get("projects") and re.search(
        r"(?im)^\s*PROJECTS\s*$", cv
    ):
        issues.append("PROJECTS section was added although the source has none.")

    return {
        "passed": not issues,
        "issues": list(dict.fromkeys(issues)),
        "missing_source_information": missing_source,
        "unsupported_target_keywords": unsupported_keywords,
    }


def ats_format_check(cv: str) -> Dict[str, Any]:
    """Deterministically assess the ATS-friendly plain-text structure."""
    headings = [
        "PROFESSIONAL SUMMARY", "SKILLS", "PROFESSIONAL EXPERIENCE",
        "EDUCATION", "ACHIEVEMENTS", "PROJECTS",
    ]
    present = [
        h for h in headings
        if re.search(rf"(?im)^\s*{re.escape(h)}\s*$", cv)
    ]
    bullet_count = len(re.findall(r"(?m)^\s*[•*-]\s+\S+", cv))
    table_like = bool(re.search(r"(?m)^\s*\|.+\|", cv))
    code_fence = "```" in cv

    score = 100
    if not present:
        score -= 30
    if bullet_count == 0 and len(cv.splitlines()) > 8:
        score -= 15
    if table_like:
        score -= 25
    if code_fence:
        score -= 20

    return {
        "format_score": max(0, min(100, score)),
        "standard_headings": present,
        "bullet_points": bullet_count,
        "table_like_content": table_like,
        "markdown_code_fence": code_fence,
    }


def apply_deterministic_review_guards(
    review: Dict[str, Any],
    cv: str,
    raw: str,
    profile: Dict[str, Any],
    job_description: Dict[str, Any],
    candidate_evidence: str = "",
) -> Dict[str, Any]:
    """Combine the independent Llama review with deterministic safety checks."""
    result = normalize_review(review)
    grounding = deterministic_grounding_check(
        cv, raw, profile, job_description, candidate_evidence
    )
    ats = ats_format_check(cv)

    result["format_score"] = min(result["format_score"], ats["format_score"])
    result["source_information_missing"] = grounding["missing_source_information"]
    result["unsupported_target_keywords"] = grounding["unsupported_target_keywords"]
    result["ats_format"] = ats

    if grounding["issues"]:
        result["issues"] = list(dict.fromkeys(
            result["issues"] + grounding["issues"]
        ))
        result["factual_consistency"] = min(
            result["factual_consistency"], 50
        )
        result["recommendation"] = "REGENERATE"

    if grounding["unsupported_target_keywords"]:
        result["recommendation"] = "REGENERATE"

    if result["factual_consistency"] < 90 or result["format_score"] < 80:
        result["recommendation"] = "REGENERATE"

    return result


"""
# Simple source-coverage guard
"""

def source_coverage_check(cv: str, profile: Dict[str, Any]) -> List[str]:
    """Backward-compatible source coverage API."""
    return deterministic_grounding_check(cv, "", profile, {}).get("issues", [])


"""
# ATS / target-keyword feedback
"""

def target_keyword_check(
    cv: str,
    job_description: Dict[str, Any],
    source_text: str = "",
    candidate_evidence: str = "",
) -> Dict[str, Any]:
    """Calculate truthful ATS keyword coverage and identify unsupported JD terms."""
    keywords = clean_list(job_description.get("keywords"))
    low = normalized_alnum(cv)
    matched = [k for k in keywords if normalized_alnum(k) in low]
    missing = [k for k in keywords if normalized_alnum(k) not in low]

    source_supported = []
    source_unsupported = []
    evidence = "\n".join(x for x in (source_text, candidate_evidence) if x)
    for keyword in keywords:
        if evidence:
            if phrase_present(evidence, keyword):
                source_supported.append(keyword)
            else:
                source_unsupported.append(keyword)

    supported_matched = [k for k in matched if k in source_supported]
    unsupported_matched = [k for k in matched if k in source_unsupported]

    eligible = source_supported
    coverage = (
        round(100 * len(supported_matched) / len(eligible))
        if eligible else 0
    )

    return {
        "target_keywords": keywords,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "keyword_coverage": coverage,
        "source_supported_keywords": source_supported,
        "source_unsupported_keywords": source_unsupported,
        "supported_matched_keywords": supported_matched,
        "unsupported_matched_keywords": unsupported_matched,
    }


def looks_like_complete_resume(text: str) -> bool:
    """Return True when the user appears to have pasted a complete CV."""
    headings = {
        "PROFESSIONAL SUMMARY", "SUMMARY", "SKILLS", "PROFESSIONAL EXPERIENCE",
        "EXPERIENCE", "EDUCATION", "ACHIEVEMENTS", "PROJECTS",
    }
    found = 0
    for line in text.splitlines():
        if line.strip().upper().rstrip(":") in headings:
            found += 1
    return found >= 2 or len(text.splitlines()) >= 12


def deterministic_user_additions(base_cv: str, edited_text: str) -> List[str]:
    """Return user-added/replaced lines without asking an LLM to detect them."""
    base_lines = [re.sub(r"\s+", " ", x.strip()) for x in base_cv.splitlines() if x.strip()]
    edited_lines = [re.sub(r"\s+", " ", x.strip()) for x in edited_text.splitlines() if x.strip()]
    base_norm = {normalized_alnum(x) for x in base_lines if normalized_alnum(x)}
    additions = []
    matcher = difflib.SequenceMatcher(
        None,
        [normalized_alnum(x) for x in base_lines],
        [normalized_alnum(x) for x in edited_lines],
        autojunk=False,
    )
    for tag, _, _, j1, j2 in matcher.get_opcodes():
        if tag in {"insert", "replace"}:
            for line in edited_lines[j1:j2]:
                if normalized_alnum(line) and normalized_alnum(line) not in base_norm:
                    if line.strip().upper().rstrip(":") in {
                        "PROFESSIONAL SUMMARY", "SUMMARY", "SKILLS",
                        "PROFESSIONAL EXPERIENCE", "EXPERIENCE", "EDUCATION",
                        "ACHIEVEMENTS", "PROJECTS",
                    }:
                        continue
                    if line not in additions:
                        additions.append(line)
    return additions


def classify_user_addition(line: str) -> str:
    """Classify a free-form user addition conservatively using explicit wording cues."""
    text = clean_string(line)
    low = text.lower()
    if re.match(r"(?i)^(skills?|technologies?)\s*:", text):
        return "skills"
    if re.match(r"(?i)^(project|project name|projects?)\s*:", text) or "project name is" in low:
        return "projects"
    if re.match(r"(?i)^(achievement|achievements?)\s*:", text):
        return "achievements"
    if re.match(r"(?i)^(education|degree)\s*:", text):
        return "education"
    if re.match(r"(?i)^(experience|employment|current employer|previous)\s*:", text):
        return "experience"
    if re.match(r"(?i)^(summary|profile)\s*:", text):
        return "summary"
    if re.search(r"\b(project|application|dashboard|system|model|solution)\b", low):
        return "projects"
    if re.search(r"\b(reduced|increased|improved|saved|cut|achieved|delivered|generated|decreased|grew|raised)\b", low):
        return "achievements"
    return "other"


def merge_additions_into_cv(base_cv: str, additions: List[str]) -> str:
    """Merge additions-only input into the existing CV without discarding the CV."""
    if not additions:
        return base_cv
    sections: Dict[str, List[str]] = {}
    other: List[str] = []
    for line in additions:
        field = classify_user_addition(line)
        if field == "other":
            other.append(line)
        else:
            sections.setdefault(field, []).append(line)

    result = base_cv.rstrip()
    section_titles = {
        "summary": "PROFESSIONAL SUMMARY",
        "skills": "SKILLS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "achievements": "ACHIEVEMENTS",
        "projects": "PROJECTS",
    }

    def append_to_section(text: str, heading: str, values: List[str]) -> str:
        lines = text.splitlines()
        heading_index = next(
            (i for i, x in enumerate(lines) if x.strip().upper().rstrip(":") == heading),
            None,
        )
        if heading_index is None:
            block = ["", heading] + [f"• {v}" for v in values]
            return text.rstrip() + "\n" + "\n".join(block)
        next_heading = len(lines)
        known = set(section_titles.values())
        for i in range(heading_index + 1, len(lines)):
            if lines[i].strip().upper().rstrip(":") in known:
                next_heading = i
                break
        insertion = [f"• {v}" for v in values]
        lines[next_heading:next_heading] = insertion + ([""] if next_heading < len(lines) else [])
        return "\n".join(lines).rstrip()

    for field, values in sections.items():
        result = append_to_section(result, section_titles[field], values)
    if other:
        # Preserve unclassified user text rather than dropping it.
        result = result.rstrip() + "\n\n" + "\n".join(f"• {x}" for x in other)
    return result.strip()


def profile_additions_from_text(additions: List[str]) -> Dict[str, Any]:
    """Build candidate evidence from explicit user additions without LLM classification."""
    result: Dict[str, Any] = {
        "name": "",
        "contact": {"email": "", "phone": "", "location": ""},
        "summary": "",
        "education": [],
        "experience": [],
        "skills": [],
        "achievements": [],
        "projects": [],
    }
    for line in additions:
        field = classify_user_addition(line)
        text = re.sub(r"(?i)^(skills?|technologies?|project|project name|projects?|achievement|achievements?|education|degree|experience|employment|current employer|previous|summary|profile)\s*:\s*", "", clean_string(line)).strip()
        if field == "skills":
            result["skills"].extend(clean_list([x.strip() for x in text.split(",") if x.strip()]))
        elif field in {"education", "experience", "achievements", "projects"}:
            result[field].append(text or clean_string(line))
        elif field == "summary":
            result["summary"] = text or clean_string(line)
        elif field == "other":
            # Keep unclassified text as a project/achievement candidate only when the
            # wording itself contains a strong project/achievement cue; otherwise it
            # remains in the CV and is not promoted into structured facts.
            continue
    return normalize_profile(result)


def user_review_and_revision(
    cv: str,
    profile: Dict[str, Any],
    raw: str,
    job_description: Dict[str, Any],
    model: str,
    host: str,
    timeout: int,
) -> tuple[str, Dict[str, Any]]:
    """Allow optional user editing, persist explicit user additions, and refine with Gemma."""
    # Display the actual generated resume BEFORE asking the user to edit it.
    print("\n" + "=" * 70)
    print("[USER REVIEW] GENERATED RESUME")
    print("=" * 70)
    print(cv)
    print("=" * 70)
    print("[USER REVIEW] Generated resume is ready for review.")
    print("[USER REVIEW] Type Y to edit/refine it, or press Enter to keep it.")
    try:
        choice = input("[USER REVIEW] Edit resume? (Y/N): ").strip().lower()
    except EOFError:
        return cv, profile
    if choice != "y":
        return cv, profile

    print("[USER REVIEW] Paste your edited version below.")
    print("[USER REVIEW] You may copy the generated resume shown above, make changes,")
    print("[USER REVIEW] and paste the complete edited resume here.")
    print("[USER REVIEW] Type END on a new line when finished.")
    lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip().upper() == "END":
            break
        lines.append(line)
    edited = "\n".join(lines).strip()
    if not edited:
        return cv, profile

    # ---------------------------------------------------------
    # Deterministic user-change detection.
    # The user's input is never sent to an LLM merely to decide what was added.
    # ---------------------------------------------------------
    if looks_like_complete_resume(edited):
        user_addition_lines = deterministic_user_additions(cv, edited)
        edited_for_refinement = edited
        print("[USER REVIEW] Input mode: COMPLETE RESUME")
    else:
        user_addition_lines = [x.strip() for x in edited.splitlines() if x.strip()]
        edited_for_refinement = merge_additions_into_cv(cv, user_addition_lines)
        print("[USER REVIEW] Input mode: ADDITIONS ONLY")

    print(f"[USER REVIEW] Deterministic user changes detected: {len(user_addition_lines)}")
    if user_addition_lines:
        print("[USER REVIEW] Accepted user-provided changes:")
        for item in user_addition_lines:
            print(f"  + {item}")

    user_additions = profile_additions_from_text(user_addition_lines)
    # Keep the exact user-added text available to the outer pipeline so ATS and
    # grounding checks can treat explicit user additions as candidate evidence.
    profile["_user_additions_evidence"] = "\n".join(user_addition_lines)

    # Merge explicit additions into the existing profile. Existing source facts
    # are preserved; only non-empty user-provided values are added/replaced.
    for field in ("education", "experience", "skills", "achievements", "projects"):
        additions = clean_list(user_additions.get(field))
        if additions:
            existing = clean_list(profile.get(field))
            for item in additions:
                if item not in existing:
                    existing.append(item)
            profile[field] = existing

    for field in ("name", "summary"):
        value = clean_string(user_additions.get(field))
        if value:
            profile[field] = value

    additions_contact = user_additions.get("contact")
    if isinstance(additions_contact, dict):
        contact = profile.get("contact")
        if not isinstance(contact, dict):
            contact = {}
        for field in ("email", "phone", "location"):
            value = clean_string(additions_contact.get(field))
            if value:
                contact[field] = value
        profile["contact"] = contact

    # ATS feedback on the user's edited/merged resume before LLM refinement.
    edited_ats = target_keyword_check(
        edited_for_refinement,
        job_description,
        raw,
        candidate_evidence="\n".join(user_addition_lines),
    )

    print("\n" + "=" * 70)
    print("[USER REVIEW] ATS FEEDBACK FOR EDITED RESUME")
    print("=" * 70)
    print(f"Keyword coverage: {edited_ats['keyword_coverage']}%")

    if edited_ats["matched_keywords"]:
        print("Matched keywords:", ", ".join(edited_ats["matched_keywords"]))
    else:
        print("Matched keywords: None")

    if edited_ats["missing_keywords"]:
        print("Missing keywords:", ", ".join(edited_ats["missing_keywords"]))
    else:
        print("Missing keywords: None")

    if edited_ats["source_unsupported_keywords"]:
        print(
            "Unsupported JD keywords (must NOT be added unless source-supported):",
            ", ".join(edited_ats["source_unsupported_keywords"]),
        )
    print("=" * 70)

    # ---------------------------------------------------------
    # EXISTING LLM refinement pipeline.
    # The only addition is that the updated profile is now supplied to it.
    # ---------------------------------------------------------
    prompt = f"""You are refining a user's edited resume.

Return the refined resume as plain text only.

USER REVIEW AND ITERATIVE REVISION RULES:

1. Preserve all candidate facts from SOURCE PROFILE.

2. Information explicitly added or corrected by the user in USER-EDITED RESUME
   is valid candidate information and must be preserved.

3. Refine the user's edited resume for wording, organization,
   ATS readability, and alignment with the TARGET JOB DESCRIPTION.

4. Identify missing JD keywords that are explicitly supported by
   the SOURCE PROFILE and incorporate those keywords naturally
   where appropriate.

5. If an achievement exists in SOURCE PROFILE but is missing from
   the USER-EDITED RESUME, it may be restored using the source wording
   or a faithful professional rewrite.

6. Do NOT create an achievement when SOURCE PROFILE contains none.

7. Do NOT add a JD keyword merely because it appears in the job description.

8. A JD skill, technology, responsibility, qualification, or keyword
   may be added to the resume ONLY when the UPDATED SOURCE PROFILE,
   ORIGINAL SOURCE TEXT, or explicit USER-EDITED RESUME content supports it.

9. Never infer candidate experience from the JD.

10. Never convert a transferable skill, similarity, interest, assumption,
    or potential into actual professional experience.

11. Never invent skills, achievements, metrics, employers, job titles,
    dates, degrees, certifications, projects, technologies,
    responsibilities, or experience.

12. Do not claim Python, SQL, AI, Machine Learning, LLM, RAG,
    LangChain, LlamaIndex, Hugging Face, PyTorch, TensorFlow,
    vector databases, or similar technologies unless explicitly
    supported by the UPDATED SOURCE PROFILE or original source.

13. If a JD keyword is not supported by the source or explicit user addition,
    leave it out of the resume. It should remain a missing keyword in ATS feedback.

14. Preserve the meaning of the original candidate information.

15. Do not include JD analysis, ATS analysis, keyword lists,
    recommendations, or explanations inside the resume.

16. Keep optional sections only when they contain candidate data.

UPDATED SOURCE PROFILE:
{json.dumps(profile, indent=2, ensure_ascii=False)}

TARGET JOB DESCRIPTION:
{json.dumps(job_description, indent=2, ensure_ascii=False)}

ORIGINAL SOURCE TEXT:
{raw}

USER-EDITED RESUME:
{edited_for_refinement}

"""

    try:
        refined = call_ollama(
            model,
            prompt,
            host,
            timeout=timeout,
            json_mode=False,
            num_predict=2000,
        )

        refined = clean_cv_text(refined)

        if user_addition_lines and not user_additions_preserved(refined, user_addition_lines):
            print("[WARN] LLM refinement dropped or materially changed explicit user content.")
            print("[WARN] Preserving the user's edited resume instead of losing user changes.")
            refined = edited_for_refinement

        # ATS feedback after LLM refinement.
        refined_ats = target_keyword_check(
            refined,
            job_description,
            raw,
            candidate_evidence="\n".join(user_addition_lines),
        )

        print("\n" + "=" * 70)
        print("[USER REVIEW] ATS FEEDBACK AFTER LLM REFINEMENT")
        print("=" * 70)
        print(
            f"Keyword coverage: "
            f"{refined_ats['keyword_coverage']}%"
        )

        if refined_ats["matched_keywords"]:
            print(
                "Matched keywords:",
                ", ".join(refined_ats["matched_keywords"]),
            )
        else:
            print("Matched keywords: None")

        if refined_ats["missing_keywords"]:
            print(
                "Missing keywords:",
                ", ".join(refined_ats["missing_keywords"]),
            )
        else:
            print("Missing keywords: None")

        if refined_ats["unsupported_matched_keywords"]:
            print(
                "WARNING - unsupported JD keywords found in CV:",
                ", ".join(
                    refined_ats["unsupported_matched_keywords"]
                ),
            )

        print("=" * 70)

        return refined, profile

    except Exception as exc:
        print(f"[WARN] User revision fallback: {exc}")
        return edited_for_refinement, profile


# ---------------------------------------------------------
# User-addition preservation guard
# ---------------------------------------------------------

def user_additions_preserved(cv: str, additions: List[str]) -> bool:
    """Check that explicit user additions remain represented after refinement."""
    cv_normalized = normalized_alnum(cv)
    for item in additions:
        item = clean_string(item)
        if not item:
            continue
        if phrase_present(cv, item):
            continue
        tokens = [t for t in normalized_alnum(item).split() if len(t) >= 4]
        if tokens and sum(t in cv_normalized for t in tokens) / len(tokens) >= 0.60:
            continue
        return False
    return True


"""
# DOCX
"""

SECTION_NAMES = {
    "SUMMARY": "PROFESSIONAL SUMMARY",
    "PROFESSIONAL SUMMARY": "PROFESSIONAL SUMMARY",
    "SKILLS": "SKILLS",
    "EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "PROFESSIONAL EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "EDUCATION": "EDUCATION",
    "ACHIEVEMENTS": "ACHIEVEMENTS",
    "PROJECTS": "PROJECTS",
}


def save_docx(text: str, path: Path) -> None:
    """Create a clean ATS-friendly DOCX from the generated plain text."""
    if Document is None:
        raise RuntimeError("Install python-docx first: pip install python-docx")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Pt(45)
    section.left_margin = section.right_margin = Pt(50)

    lines = [x.strip() for x in text.splitlines()]

    for index, line in enumerate(lines):
        if not line:
            continue

        normalized = re.sub(r"^[•*-]\s*", "", line).strip()
        key = normalized.upper().rstrip(":")

        if index == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(normalized)
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(18)
            continue

        if key in SECTION_NAMES:
            p = doc.add_paragraph()
            r = p.add_run(SECTION_NAMES[key])
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(11.5)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            continue

        p = doc.add_paragraph()
        if line.startswith("•") or re.match(r"^[*-]\s+", line):
            p.style = doc.styles["List Bullet"]
            content = re.sub(r"^[•*-]\s*", "", line)
        else:
            content = normalized

        r = p.add_run(content)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)

    doc.save(path)


"""
# Main pipeline
"""

def process_profile(
    file_path: Path,
    output_dir: Path,
    args: argparse.Namespace,
    job_description: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    print(f"\n[INFO] Processing {file_path.name}")

    raw = extract_text_from_file(file_path)

    print("[1/6] Extracting source facts with Gemma...")
    profile = extract_profile(
        raw,
        args.gemma_model,
        args.ollama_host,
        args.timeout,
    )

    print("[2/6] Parsing job description with Gemma...")
    job_description = job_description or {}
    save_job_description(job_description, output_dir)

    print("[3/6] Matching candidate facts to job keywords...")
    keyword_feedback = target_keyword_check("", job_description, raw)

    print("[4/6] Generating tailored ATS-friendly CV with Gemma...")
    cv = generate_cv(
        profile,
        raw,
        args.gemma_model,
        args.ollama_host,
        args.timeout,
        job_description=job_description,
    )

    keyword_feedback = target_keyword_check(cv, job_description, raw)
    print(
        f"[ATS] Target keyword coverage: {keyword_feedback['keyword_coverage']}% "
        f"({len(keyword_feedback['matched_keywords'])}/{len(keyword_feedback['target_keywords'])})"
    )
    if keyword_feedback["missing_keywords"]:
        print("[ATS] Missing target keywords:", ", ".join(keyword_feedback["missing_keywords"]))

    print("[5/6] Reviewing CV independently with Llama...")
    review = review_cv(
        cv,
        profile,
        raw,
        args.llama_model,
        args.ollama_host,
        args.timeout,
        job_description=job_description,
    )

    review = apply_deterministic_review_guards(
        review, cv, raw, profile, job_description
    )

    # Target-job keyword coverage is part of the ATS feedback.
    review["target_keyword_coverage"] = keyword_feedback["keyword_coverage"]
    review["target_keywords"] = keyword_feedback["target_keywords"]
    review["matched_target_keywords"] = keyword_feedback["matched_keywords"]
    review["missing_target_keywords"] = keyword_feedback["missing_keywords"]

    # One and only one controlled regeneration, preserving the original behavior.
    if review.get("recommendation") == "REGENERATE":
        reasons = "\n".join(
            review.get("issues", [])[:8]
            + review.get("missing_keywords", [])[:8]
            + keyword_feedback.get("missing_keywords", [])[:8]
        )

        print("[6/6] Review requested regeneration; retrying Gemma once...")
        regenerated = generate_cv(
            profile,
            raw,
            args.gemma_model,
            args.ollama_host,
            args.timeout,
            regeneration_reason=reasons,
            job_description=job_description,
        )
        if regenerated.strip():
            cv = regenerated

        print("[INFO] Running final Llama review...")
        review = review_cv(
            cv,
            profile,
            raw,
            args.llama_model,
            args.ollama_host,
            args.timeout,
            job_description=job_description,
        )
        keyword_feedback = target_keyword_check(cv, job_description, raw)
        review["target_keyword_coverage"] = keyword_feedback["keyword_coverage"]
        review["target_keywords"] = keyword_feedback["target_keywords"]
        review["matched_target_keywords"] = keyword_feedback["matched_keywords"]
        review["missing_target_keywords"] = keyword_feedback["missing_keywords"]

        review = apply_deterministic_review_guards(
            review, cv, raw, profile, job_description
        )

    # Optional user review/editing. This is deliberately interactive and does
    # not create an extra draft file.
    cv, profile = user_review_and_revision(
        cv,
        profile,
        raw,
        job_description,
        args.gemma_model,
        args.ollama_host,
        args.timeout,
    )
    cv = clean_cv_text(cv)
    user_evidence = clean_string(profile.pop("_user_additions_evidence", ""))

    final_grounding = deterministic_grounding_check(
        cv, raw, profile, job_description, user_evidence
    )
    if final_grounding["issues"]:
        # User-provided additions are candidate evidence. Do not silently discard
        # the user's edited CV because an original-source guard cannot account for it.
        print("[WARN] User-edited CV has grounding warnings; preserving the user's changes.")
        print("[WARN] " + " | ".join(final_grounding["issues"][:5]))

    keyword_feedback = target_keyword_check(
        cv, job_description, raw, candidate_evidence=user_evidence
    )
    review["target_keyword_coverage"] = keyword_feedback["keyword_coverage"]
    review["target_keywords"] = keyword_feedback["target_keywords"]
    review["matched_target_keywords"] = keyword_feedback["matched_keywords"]
    review["missing_target_keywords"] = keyword_feedback["missing_keywords"]
    review["source_supported_keywords"] = keyword_feedback["source_supported_keywords"]
    review["unsupported_matched_keywords"] = keyword_feedback["unsupported_matched_keywords"]

    review = apply_deterministic_review_guards(
        review, cv, raw, profile, job_description, user_evidence
    )

    stem = file_path.stem

    # Preserve the existing per-profile structured/review outputs.
    (output_dir / f"{stem}_structured.json").write_text(
        json.dumps(profile, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    docx_path = output_dir / f"{stem}_final_cv.docx"
    save_docx(cv, docx_path)

    print(
        f"[RESULT] {file_path.name}: "
        f"status={review.get('recommendation')}, "
        f"score={review.get('score', 0)}, "
        f"target_keyword_coverage={keyword_feedback['keyword_coverage']}%"
    )

    return {
        "input": file_path.name,
        "output_cv": docx_path.name,
        "score": review.get("score", 0),
        "keyword_coverage": review.get("keyword_coverage", 0),
        "target_keyword_coverage": keyword_feedback["keyword_coverage"],
        "factual_consistency": review.get("factual_consistency", 0),
        "format_score": review.get("format_score", 0),
        "status": review.get("recommendation", "REGENERATE"),
        "issues": review.get("issues", []),
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="HPPCS01 - CV Creation using two local LLMs"
    )
    parser.add_argument("--input-dir", default=".")
    parser.add_argument("--output-dir", default=".")
    # Keep 1B as the default because this is the currently reverted working setup.
    # Use --gemma-model gemma3:4b when you intentionally want the 4B model.
    parser.add_argument("--gemma-model", default="gemma3:1b")
    parser.add_argument("--llama-model", default="llama3.2:1b")
    parser.add_argument("--ollama-host", default="http://localhost:11434")
    parser.add_argument("--timeout", type=int, default=180)
    parser.add_argument("--limit", type=int, default=10)
    parser.add_argument("--jd", default="", help="Job description text; if omitted, paste it when prompted.")
    parser.add_argument("--jd-file", default="", help="Path to a job description text file.")

    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    files = discover_input_files(input_dir)[: args.limit]
    if not files:
        raise SystemExit("No supported resume files found. Supported: PDF, DOC, DOCX.")

    jd_text = load_job_description(input_dir, args)
    if not jd_text:
        raise SystemExit("A job description is required. Provide job_description.txt, --jd-file, --jd, or paste it when prompted.")

    print("=" * 70)
    print("Capstone_Project-HPPCS01")
    print("=" * 70)
    print(f"Input directory : {input_dir.resolve()}")
    print(f"Output directory: {output_dir.resolve()}")
    print(f"Gemma model     : {args.gemma_model}")
    print(f"Review model    : {args.llama_model}")
    print(f"Timeout         : {args.timeout}s")
    print(f"Supported input : {", ".join(sorted(SUPPORTED_EXTENSIONS))}")
    print(f"Files discovered : {len(files)}")
    print("[INFO] Parsing job description with Gemma...")
    job_description = parse_job_description(
        jd_text, args.gemma_model, args.ollama_host, args.timeout
    )
    save_job_description(job_description, output_dir)
    print(f"[INFO] Structured JD saved: {(output_dir / 'job_description.json').name}")

    results = []

    for file_path in files:
        try:
            results.append(process_profile(file_path, output_dir, args, job_description))
        except Exception as exc:
            print(f"[ERROR] Failed to process {file_path.name}: {exc}")
            results.append(
                {
                    "input": file_path.name,
                    "output_cv": "",
                    "score": 0,
                    "keyword_coverage": 0,
                    "factual_consistency": 0,
                    "format_score": 0,
                    "status": "ERROR",
                    "issues": [str(exc)],
                }
            )

    summary = {
        "project": "Capstone_Project-HPPCS01",
        "models": {
            "generation": args.gemma_model,
            "review": args.llama_model,
        },
        "job_description": job_description,
        "profiles_processed": len(results),
        "results": results,
    }

    (output_dir / "evaluation_summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    print("\n" + "=" * 70)
    print(json.dumps(summary, indent=2, ensure_ascii=False))
    print("=" * 70)


if __name__ == "__main__":
    main()
